import os.path
from SimhashSimilarity import SimhashSimilarity
from CosineSimilarity import CosineSimilarity
import docx
from flask import current_app


class Document:
    def __init__(self, file_path):
        try:
            self.graphs = []
            self.file = docx.Document(os.path.join(current_app.config['UPLOAD_FOLDER'], file_path))
            self.filename = file_path
            for graph in self.file.paragraphs:
                self.graphs.append(graph.text)
            for table in self.file.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.graphs.append(cell.text)
        except FileNotFoundError:  # 文件没找到Error该如何处理呢？这肯定会500 Error
            pass

    @property
    def text(self):
        """
        :return:all texts in document,in the format of list
        """
        return self.graphs


class Documents:
    """
    计算文本之间的相似度方法还有很多，还可以再添加其他的计算方法。
    """
    def __init__(self, documents_list):
        self.documents = documents_list
        self.simhash_dict = dict()
        self.get_simhash_similarity()

    def get_simhash_similarity(self):
        for file1 in self.documents:
            for file2 in self.documents:
                if file1.filename != file2.filename:
                    # 在此处求相似度
                    simhashs = SimhashSimilarity(file1.text, file2.text)
                    key = str(zip(file1.filename,file2.filename))
                    self.simhash_dict[key] = simhashs

    @property
    def simhash_similarity(self):
        return self.simhash_dict


